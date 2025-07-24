from flask import Flask, request, send_file, render_template_string
import io
import re
from docx import Document

app = Flask(__name__)

# Padrões de dados sensíveis
PADROES_SENSIVEIS = {
    "CPF": r'\b\d{3}\.\d{3}\.\d{3}-\d{2}\b',
    "RG": r'\b\d{1,2}\.?\d{3}\.?\d{3}-?\d{1}\b',
    "EMAIL": r'\b[\w\.-]+@[\w\.-]+\.\w{2,}\b',
    "TELEFONE": r'\(?\d{2}\)?\s?\d{4,5}-\d{4}',
    "CEP": r'\b\d{5}-\d{3}\b',
    "CNPJ": r'\b\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}\b',
    "CARTAO": r'(?:\d[ -]*?){13,16}',
    "PLACA": r'\b[A-Z]{3}-?\d{1}[A-Z0-9]{1}\d{2}\b',
    "DATA": r'\b\d{2}/\d{2}/\d{4}\b'
}

HTML_FORM = """
<!doctype html>
<title>Tarjar Informações Sensíveis</title>
<h2>Envie um arquivo .docx e escolha quais dados deseja tarjar:</h2>
<form method="post" enctype="multipart/form-data">
  <input type="file" name="docxfile" accept=".docx" required><br><br>
  {% for chave in padroes %}
    <label><input type="checkbox" name="itens" value="{{ chave }}" checked> {{ chave }}</label><br>
  {% endfor %}
  <br>
  <button type="submit">Enviar e Tarjar</button>
</form>
"""

def copiar_e_tarjar(original_doc, padroes):
    novo_doc = Document()

    for par in original_doc.paragraphs:
        texto = par.text
        for nome, regex in padroes.items():
            texto = re.sub(regex, lambda m: "█" * len(m.group()), texto)

        # Adiciona o parágrafo com texto tarjado ao novo doc
        novo_doc.add_paragraph(texto)

    return novo_doc

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        arquivo = request.files['docxfile']
        selecionados = request.form.getlist("itens")

        if not arquivo or not arquivo.filename.endswith('.docx'):
            return "Arquivo inválido. Envie um .docx.", 400

        padroes_ativos = {k: v for k, v in PADROES_SENSIVEIS.items() if k in selecionados}

        original_doc = Document(arquivo)
        novo_doc = copiar_e_tarjar(original_doc, padroes_ativos)

        mem_file = io.BytesIO()
        novo_doc.save(mem_file)
        mem_file.seek(0)

        return send_file(
            mem_file,
            as_attachment=True,
            download_name="documento_tarjado.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    return render_template_string(HTML_FORM, padroes=PADROES_SENSIVEIS.keys())

if __name__ == '__main__':
    app.run(debug=True)
